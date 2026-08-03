from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "solo-public-project-journey.md"
OUTPUT = ROOT / "solo-public-project-journey.docx"

NAVY = "173B57"
GREEN = "176B58"
BLUE = "2E74B5"
INK = "243447"
MUTED = "667085"
PALE = "E8EEF5"
LIGHT = "F5F7FA"
WHITE = "FFFFFF"
RED = "9B1C1C"


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")


def add_inline(paragraph, text, size=11, color=INK):
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, "Consolas", 9.5, NAVY)
            run.font.highlight_color = None
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, size=size, color=color, bold=True)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size=size, color=color)


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, size=9, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, NAVY, 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def add_cover(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(92)
    p.paragraph_format.space_after = Pt(16)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("PRODUCT JOURNEY")
    set_run_font(r, size=10, color=GREEN, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("Solo Public Project")
    set_run_font(r, size=30, color=NAVY, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(32)
    r = p.add_run("End-to-end owner and collaborator workflow")
    set_run_font(r, size=15, color=BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.right_indent = Inches(0.18)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    p_pr = p._p.get_or_add_pPr()
    p_shd = OxmlElement("w:shd")
    p_shd.set(qn("w:fill"), PALE)
    p_pr.append(p_shd)
    add_inline(p, "Permissions • uploads • review decisions • completion • acceptance tests", 10.5, NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(130)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("DocsNDocs")
    set_run_font(r, size=12, color=NAVY, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Implementation and verification guide")
    set_run_font(r, size=9.5, color=MUTED)
    doc.add_page_break()


def add_contents(doc, headings):
    p = doc.add_paragraph("Contents", style="Heading 1")
    p.paragraph_format.space_before = Pt(0)
    for level, title in headings:
        if level > 2:
            continue
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.22 if level == 2 else 0)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(title)
        set_run_font(r, size=10.5 if level == 2 else 11, color=NAVY, bold=(level == 1))
    doc.add_page_break()


def add_markdown_table(doc, rows):
    headers = rows[0]
    body = rows[2:]
    cols = len(headers)
    table = doc.add_table(rows=1 + len(body), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    max_lens = [max(len(r[i]) if i < len(r) else 0 for r in [headers] + body) for i in range(cols)]
    total = sum(max(max_lens[i], 8) for i in range(cols))
    widths = [int(9360 * max(max_lens[i], 8) / total) for i in range(cols)]
    widths[-1] += 9360 - sum(widths)
    set_table_geometry(table, widths)
    set_repeat_table_header(table.rows[0])
    for ridx, values in enumerate([headers] + body):
        for cidx, value in enumerate(values):
            cell = table.cell(ridx, cidx)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            if ridx == 0:
                shade(cell, PALE)
            elif ridx % 2 == 0:
                shade(cell, LIGHT)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            add_inline(p, value.strip(), 9.3, NAVY if ridx == 0 else INK)
            for run in p.runs:
                if ridx == 0:
                    run.bold = True
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def build():
    raw = SOURCE.read_text(encoding="utf-8").replace("â€“", "–").replace("Ã—", "×")
    lines = raw.splitlines()
    headings = []
    for line in lines:
        m = re.match(r"^(#{2,3})\s+(.+)$", line)
        if m:
            headings.append((len(m.group(1)) - 1, m.group(2)))

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    configure_styles(doc)

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hr = hp.add_run("Solo Public Project Journey")
    set_run_font(hr, size=9, color=MUTED, bold=True)
    add_page_field(section.footer.paragraphs[0])

    add_cover(doc)
    add_contents(doc, headings)

    paragraph_buffer = []
    table_rows = []
    in_code = False
    code_lines = []

    def flush_para():
        nonlocal paragraph_buffer
        if paragraph_buffer:
            p = doc.add_paragraph()
            add_inline(p, " ".join(x.strip() for x in paragraph_buffer))
            paragraph_buffer = []

    def flush_table():
        nonlocal table_rows
        if table_rows:
            add_markdown_table(doc, table_rows)
            table_rows = []

    for line in lines[1:]:
        if line.startswith("```"):
            flush_para()
            if in_code:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.18)
                p.paragraph_format.right_indent = Inches(0.18)
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(8)
                p.paragraph_format.line_spacing = 1.0
                r = p.add_run("\n".join(code_lines))
                set_run_font(r, "Consolas", 9, NAVY)
                pPr = p._p.get_or_add_pPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:fill"), LIGHT)
                pPr.append(shd)
                code_lines = []
            in_code = not in_code
            continue
        if in_code:
            code_lines.append(line)
            continue
        if line.startswith("|"):
            flush_para()
            table_rows.append([c.strip() for c in line.strip().strip("|").split("|")])
            continue
        flush_table()
        if not line.strip():
            flush_para()
            continue
        h = re.match(r"^(#{2,4})\s+(.+)$", line)
        if h:
            flush_para()
            level = min(len(h.group(1)) - 1, 3)
            doc.add_paragraph(h.group(2), style=f"Heading {level}")
            continue
        bullet = re.match(r"^\s*-\s+(.+)$", line)
        if bullet:
            flush_para()
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, bullet.group(1))
            continue
        numbered = re.match(r"^\s*\d+\.\s+(.+)$", line)
        if numbered:
            flush_para()
            p = doc.add_paragraph(style="List Number")
            add_inline(p, numbered.group(1))
            continue
        paragraph_buffer.append(line)
    flush_para()
    flush_table()

    props = doc.core_properties
    props.title = "Solo Public Project Journey"
    props.subject = "End-to-end DocsNDocs owner and collaborator workflow"
    props.author = "DocsNDocs"
    props.keywords = "solo project, public project, collaborators, document approval, Vercel Blob"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
